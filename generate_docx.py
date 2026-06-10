import docx
from docx.shared import Pt, Inches
import re
import sys

def parse_markdown_to_docx(md_path, docx_path):
    doc = docx.Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code_block = False
    code_content = []
    
    for line in lines:
        line = line.strip('\n')
        
        if line.startswith('```'):
            if in_code_block:
                # End code block
                p = doc.add_paragraph()
                p.style = 'No Spacing'
                run = p.add_run('\n'.join(code_content))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.5)
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            continue
            
        if in_code_block:
            code_content.append(line)
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:], 0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], 1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], 2)
        elif line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. '):
            doc.add_paragraph(line[3:], style='List Number')
        elif line.startswith('2. '):
            doc.add_paragraph(line[3:], style='List Number')
        elif line.startswith('3. '):
            doc.add_paragraph(line[3:], style='List Number')
        elif line.strip() == '':
            continue
        else:
            # Handle bold text
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                else:
                    p.add_run(part)

    doc.save(docx_path)
    print(f"Successfully created {docx_path}")

md_file = r"C:\Users\aarit\.gemini\antigravity\brain\ce1d0466-484d-4db6-868b-cd7e851a2db2\Asterisk_Integration_Guide.md"
docx_file = r"C:\Users\aarit\OneDrive\Desktop\NeoX_Asterisk_Integration_Guide.docx"
parse_markdown_to_docx(md_file, docx_file)
